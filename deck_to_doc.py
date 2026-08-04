#!/usr/bin/env python3
"""Turn a deck's index.src.html into an editable Google Doc.

The Doc is the writing surface: Steven edits prose there, then Claude reads it
back and rebuilds the deck. Structure is deliberately plain — headings and
paragraphs, one block per slide — so editing feels like editing a document, not
markup. The bracketed component labels are the only scaffolding, and they exist
so an edit can be mapped back to the right element without guessing.

    python3 deck_to_doc.py                     # build .docx and upload as a Google Doc
    python3 deck_to_doc.py --local             # build the .docx only, no upload
    python3 deck_to_doc.py --name my-doc       # name the Doc (a new name = a new Doc)

Run from a deck folder (needs index.src.html alongside).
"""
import re, sys, html, pathlib, subprocess, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = pathlib.Path.cwd()
SRC = HERE / "index.src.html"
if not SRC.exists():
    raise SystemExit("no index.src.html here — run this from a deck folder")

DRIVE_DIR = 'gdrive:"WorldFirst Docs"'
raw = SRC.read_text(encoding="utf-8")

title_m = re.search(r"<title>(.*?)</title>", raw, re.S)
DECK_TITLE = html.unescape(title_m.group(1).strip()) if title_m else HERE.name


def text_of(frag: str) -> str:
    """Visible text of an HTML fragment, whitespace collapsed."""
    frag = re.sub(r"<br\s*/?>", " ", frag)
    frag = re.sub(r"<[^>]+>", "", frag)
    return html.unescape(re.sub(r"\s+", " ", frag)).strip()


def divs(frag: str, cls_re: str):
    """Inner HTML of every <div class="...cls_re..."> in frag, matching nested divs
    by depth. A non-greedy regex stops at the first </div>, which for .tbl and
    .agenda is an inner row, not the container."""
    out = []
    for m in re.finditer(r'<div class="([^"]*)"[^>]*>', frag):
        if not re.search(cls_re, m.group(1)):
            continue
        depth, i = 1, m.end()
        while depth and i < len(frag):
            nxt = re.search(r"</?div\b", frag[i:])
            if not nxt:
                break
            j = i + nxt.start()
            depth += -1 if frag[j:j + 2] == "</" else 1
            i = j + (6 if frag[j:j + 2] == "</" else 4)
        out.append(frag[m.end():i - 6])
    return out


def parse_slides(src: str):
    slides = []
    for m in re.finditer(r'<section class="slide([^"]*)">(.*?)</section>', src, re.S):
        cls, inner = m.group(1).strip(), m.group(2)
        blocks = []

        for t in re.findall(r'<h[12] class="title">(.*?)</h[12]>', inner, re.S):
            blocks.append(("Title", [text_of(t)]))
        for t in re.findall(r'<p class="sub">(.*?)</p>', inner, re.S):
            blocks.append(("Subheading", [text_of(t)]))

        for b in divs(inner, r'\bbody\b'):
            items = re.findall(r"<li>(.*?)</li>", b, re.S)
            if items:
                blocks.append(("Bullets", [text_of(i) for i in items]))
            else:
                paras = re.findall(r"<p>(.*?)</p>", b, re.S)
                blocks.append(("Body", [text_of(p) for p in paras] or [text_of(b)]))

        for chunk in divs(inner, r'\bagenda\b'):
            rows = re.findall(r"<b>(.*?)</b>\s*<span>(.*?)</span>", chunk, re.S)
            if rows:
                blocks.append(("Contents", [f"{text_of(x)} — {text_of(y)}" for x, y in rows]))

        for c in divs(inner, r'\bcards\b'):
            rows = re.findall(r"<h3>(.*?)</h3>(?:\s*<p>(.*?)</p>)?", c, re.S)
            if rows:
                blocks.append(("Cards", [f"{text_of(h)}" + (f" — {text_of(p)}" if p else "")
                                         for h, p in rows]))

        for tb in divs(inner, r'\btbl\b'):
            rows = []
            for tr in re.finditer(r'<div class="tr([^"]*)">(.*?)</div>', tb, re.S):
                cells = [text_of(c) for c in re.findall(r"<span[^>]*>(.*?)</span>", tr.group(2), re.S)]
                if cells:
                    rows.append(("   ".join(cells), "th" in tr.group(1)))
            if rows:
                blocks.append(("Table", rows))

        for rf in divs(inner, r'\brefs\b'):
            rows = re.findall(r"<h3>(.*?)</h3>\s*<p>(.*?)</p>", rf, re.S)
            if rows:
                blocks.append(("References", [f"{text_of(a)} — {text_of(b)}" for a, b in rows]))

        for st in divs(inner, r'\bstats\b'):
            rows = re.findall(r"<b>(.*?)</b>\s*<span>(.*?)</span>", st, re.S)
            if rows:
                blocks.append(("Figures", [f"{text_of(a)} — {text_of(b)}" for a, b in rows]))

        for tl in divs(inner, r'\btimeline\b'):
            heads = [text_of(x) for x in re.findall(r"<div>(.*?)</div>", tl, re.S)][:5]
            bars = [text_of(x) for x in re.findall(r'<div class="bar[^"]*"[^>]*>(.*?)</div>', tl, re.S)]
            if heads or bars:
                blocks.append(("Timeline", [" · ".join(heads)] + bars))

        for q in re.findall(r'<p class="quote[^"]*">(.*?)</p>', inner, re.S):
            blocks.append(("Pull quote", [text_of(q)]))
        for q in re.findall(r'<div class="quotebox"[^>]*>(.*?)</div>', inner, re.S):
            blocks.append(("Quoted block", [text_of(q)]))
        for n in re.findall(r'<p class="note">(.*?)</p>', inner, re.S):
            blocks.append(("Footnote", [text_of(n)]))

        links = re.findall(r'href="([^"]+)"', inner)
        slides.append({"cls": cls, "blocks": blocks, "links": links})
    return slides


slides = parse_slides(raw)

# ---------- document ------------------------------------------------------
INK = RGBColor(0x1F, 0x23, 0x23)
GREY = RGBColor(0x6E, 0x74, 0x78)
PINK = RGBColor(0xC8, 0x10, 0x4C)

doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(1.0)
base = doc.styles["Normal"]
base.font.name = "Helvetica Neue"
base.font.size = Pt(10.5)
base.font.color.rgb = INK
base.paragraph_format.space_after = Pt(5)


def para(text="", size=10.5, bold=False, color=INK, before=0, after=5, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return p


para(DECK_TITLE, size=19, bold=True, after=2)
para("Editable copy of every slide. Change the wording here, then tell Claude to "
     "rebuild the deck from this Doc. Rough notes are fine — Claude tidies the register.",
     size=9.5, color=GREY, after=3)
para(f"{len(slides)} slides", size=9.5, color=GREY, after=16)

for n, s in enumerate(slides, 1):
    bg = next((c for c in s["cls"].split() if c.startswith("bg-")), "")
    head = doc.add_paragraph()
    head.paragraph_format.space_before = Pt(15)
    head.paragraph_format.space_after = Pt(4)
    r = head.add_run(f"Slide {n}")
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = PINK
    if bg:
        r2 = head.add_run(f"   {bg}")
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GREY

    for label, content in s["blocks"]:
        para(label, size=8.5, bold=True, color=GREY, before=5, after=2)
        if label == "Table":
            for row, is_head in content:
                para(row, size=9.5, color=GREY if is_head else INK, after=1)
        elif label in ("Bullets", "Contents", "Cards", "Figures", "Timeline", "References"):
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(item)
                rr.font.size = Pt(9.5)
                rr.font.color.rgb = INK
        else:
            for item in content:
                para(item, size=10.5, after=3)

    for l in s["links"]:
        para(l, size=8.5, color=GREY, italic=True, after=1)

name = HERE.name + "-content"
if "--name" in sys.argv:
    name = sys.argv[sys.argv.index("--name") + 1]
out = HERE / f"{name}.docx"
doc.save(out)
print(f"built {out.name} — {len(slides)} slides")

if "--local" in sys.argv:
    raise SystemExit(0)

if not shutil.which("rclone"):
    raise SystemExit("rclone not found; .docx left in place")

cmd = (f'rclone copy "{out}" {DRIVE_DIR} '
       f'--drive-import-formats docx --drive-allow-import-name-change')
r = subprocess.run(cmd, shell=True, capture_output=True, text=True)
if r.returncode:
    raise SystemExit(f"upload failed:\n{r.stdout}{r.stderr}")
print(f"uploaded to {DRIVE_DIR} as a Google Doc")
out.unlink()
print("local .docx removed")
